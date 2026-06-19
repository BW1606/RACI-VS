from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import hierarchy_path, R_SUBCATEGORIES

ROLE_LABELS = {
    "R": "Responsible",
    "A": "Accountable",
    "C": "Consulted",
    "I": "Informed",
    "V": "Verificator",
    "S": "Signator",
}

R_SUBCATEGORY_HEADINGS = {
    "Ausführende Tätigkeit": "Ausführen",
    "Gewährleistung": "Gewährleisten",
    "Koordination": "Koordinieren",
    "Veranlassung": "Veranlassen",
    "Mitwirkung": "Mitwirken",
}

R_SUBCATEGORY_ORDER = list(R_SUBCATEGORIES)


def _add_field(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value or "—")


def _add_placeholder(doc: Document) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Details ergänzen")
    run.italic = True


def _set_table_header_style(row) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        cell._tc.get_or_add_tcPr()


def generate_function_description(fn) -> BytesIO:
    doc = Document()

    # Title block
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Funktionsbeschreibung")
    title_run.bold = True
    title_run.font.size = Pt(28)

    doc.add_heading(fn.name, level=1)

    if fn.description:
        doc.add_paragraph(fn.description)

    # 1. Organisatorische Einordnung
    doc.add_heading("1. Organisatorische Einordnung", level=2)

    parent_chain = hierarchy_path(fn.parent) or "—"
    _add_field(doc, "Übergeordnete Funktion", parent_chain)

    children_names = ", ".join(c.name for c in fn.children) if fn.children else "—"
    _add_field(doc, "Untergeordnete Funktion", children_names)

    # 2. Ziel der Funktion
    doc.add_heading("2. Ziel der Funktion", level=2)
    doc.add_paragraph(fn.purpose or "—")

    # Group task_roles by role
    by_role: dict[str, list] = {"A": [], "R": [], "C": [], "I": [], "V": [], "S": []}
    for tr in fn.task_roles:
        if tr.role in by_role:
            by_role[tr.role].append(tr)

    # 3. Ergebnisverantwortung (A)
    doc.add_heading("3. Ergebnisverantwortung (A)", level=2)
    if by_role["A"]:
        for tr in by_role["A"]:
            doc.add_paragraph(tr.task.title, style="List Bullet")
    else:
        _add_placeholder(doc)

    # 4. Durchführungsverantwortung (R)
    doc.add_heading("4. Durchführungsverantwortung (R)", level=2)
    if by_role["R"]:
        by_sub: dict[str, list] = {s: [] for s in R_SUBCATEGORY_ORDER}
        for tr in by_role["R"]:
            sub = tr.r_subcategory or "Ausführende Tätigkeit"
            if sub in by_sub:
                by_sub[sub].append(tr)
        for sub_key in R_SUBCATEGORY_ORDER:
            tasks = by_sub[sub_key]
            heading = R_SUBCATEGORY_HEADINGS[sub_key]
            p = doc.add_paragraph()
            p.add_run(heading).bold = True
            if tasks:
                for tr in tasks:
                    doc.add_paragraph(tr.task.title, style="List Bullet")
            else:
                _add_placeholder(doc)
    else:
        for sub_key in R_SUBCATEGORY_ORDER:
            heading = R_SUBCATEGORY_HEADINGS[sub_key]
            p = doc.add_paragraph()
            p.add_run(heading).bold = True
            _add_placeholder(doc)

    # 5. Beratungsverantwortung (C)
    doc.add_heading("5. Beratungsverantwortung (C)", level=2)
    if by_role["C"]:
        for tr in by_role["C"]:
            doc.add_paragraph(tr.task.title, style="List Bullet")
    else:
        _add_placeholder(doc)

    # 6. Informationsverantwortung (I)
    doc.add_heading("6. Informationsverantwortung (I)", level=2)
    if by_role["I"]:
        for tr in by_role["I"]:
            doc.add_paragraph(tr.task.title, style="List Bullet")
    else:
        _add_placeholder(doc)

    # 7. Verifikationsverantwortung (V)
    doc.add_heading("7. Verifikationsverantwortung (V)", level=2)
    if by_role["V"]:
        for tr in by_role["V"]:
            doc.add_paragraph(tr.task.title, style="List Bullet")
    else:
        _add_placeholder(doc)

    # 8. Signaturverantwortung (S)
    doc.add_heading("8. Signaturverantwortung (S)", level=2)
    if by_role["S"]:
        for tr in by_role["S"]:
            doc.add_paragraph(tr.task.title, style="List Bullet")
    else:
        _add_placeholder(doc)

    # 9. Befugnisse
    doc.add_heading("9. Befugnisse", level=2)
    doc.add_paragraph(fn.befugnisse or "—")

    # 10. Vertretung
    doc.add_heading("10. Vertretung", level=2)
    rep_name = fn.emergency_rep.name if fn.emergency_rep else "Details ergänzen"
    _add_field(doc, "Vertretung von", rep_name)
    p = doc.add_paragraph()
    p.add_run("Wird vertreten: ").bold = True
    p.add_run("Details ergänzen").italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_task_description(task) -> BytesIO:
    doc = Document()

    doc.add_heading(task.title, level=1)
    doc.add_heading("Task Details", level=2)
    _add_field(doc, "Description", task.description)

    doc.add_heading("Involved Functions", level=2)

    if task.function_roles:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Function"
        hdr[1].text = "Role"
        hdr[2].text = "Hierarchy"
        _set_table_header_style(table.rows[0])

        for fr in task.function_roles:
            row = table.add_row().cells
            row[0].text = fr.function.name
            row[1].text = f"{fr.role} — {ROLE_LABELS.get(fr.role, fr.role)}"
            row[2].text = hierarchy_path(fr.function)
    else:
        doc.add_paragraph("No functions assigned to this task.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_interface_description(fn1, fn2, shared_tasks: list[dict]) -> BytesIO:
    doc = Document()

    doc.add_heading(f"Interface: {fn1.name} ↔ {fn2.name}", level=1)

    p = doc.add_paragraph()
    p.add_run("This document describes the shared tasks between ")
    p.add_run(fn1.name).bold = True
    p.add_run(" and ")
    p.add_run(fn2.name).bold = True
    p.add_run(" and their respective roles.")

    doc.add_heading("Shared Tasks", level=2)

    if shared_tasks:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Task"
        hdr[1].text = f"Role of {fn1.name}"
        hdr[2].text = f"Role of {fn2.name}"
        hdr[3].text = "Task Description"
        _set_table_header_style(table.rows[0])

        for item in shared_tasks:
            row = table.add_row().cells
            row[0].text = item["task_title"]
            role1 = item["role_f1"]
            role2 = item["role_f2"]
            row[1].text = f"{role1} — {ROLE_LABELS.get(role1, role1)}"
            row[2].text = f"{role2} — {ROLE_LABELS.get(role2, role2)}"
            row[3].text = item["task_description"]
    else:
        doc.add_paragraph(f"{fn1.name} and {fn2.name} share no tasks.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
